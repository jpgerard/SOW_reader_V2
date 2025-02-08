"""
SOW Processor using spaCy for improved requirement extraction.
"""

import re
import logging
from pathlib import Path
from typing import List, Dict, Tuple, Optional
import spacy
import pdfplumber
from docx import Document
import nltk

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')
logger = logging.getLogger(__name__)

# Silence pdfminer and pdfplumber debug logs
logging.getLogger("pdfminer").setLevel(logging.WARNING)
logging.getLogger("pdfplumber").setLevel(logging.WARNING)

# Load spaCy model
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    import subprocess
    subprocess.check_call(["python", "-m", "spacy", "download", "en_core_web_sm"])
    nlp = spacy.load("en_core_web_sm")

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

class SOWProcessor:
    """
    Processes Statement of Work documents using spaCy NLP for improved extraction.
    """
    def __init__(self):
        self.mandatory_keywords = [
            "shall", "must", "required to", "responsible for", "directed to", "required"
        ]
        self.informative_keywords = [
            "will", "plans to", "intends to", "should", "expected to", "recommended"
        ]
        self.action_verbs = [
            "provide", "implement", "support", "develop", "maintain",
            "ensure", "perform", "comply", "submit", "deliver"
        ]
        
        # Precompile regex patterns
        self.mandatory_patterns = [
            re.compile(rf'\b{re.escape(kw)}\b', re.IGNORECASE) 
            for kw in self.mandatory_keywords
        ]
        self.informative_patterns = [
            re.compile(rf'\b{re.escape(kw)}\b', re.IGNORECASE) 
            for kw in self.informative_keywords
        ]
        
        # Section header patterns
        self.header_patterns = [
            re.compile(r'^([A-Z](?:\.\d+)?|\d+\.\d+(?:\.\d+)?)\s+(.+)$'),
            re.compile(r'^Section\s+(\d+)[\.:]?\s*(.+)$', re.IGNORECASE)
        ]

    def _join_lines(self, lines: List[str]) -> Tuple[str, List[int]]:
        """Joins lines with newline characters and returns line offsets."""
        line_offsets = []
        current_offset = 0
        for line in lines:
            line_offsets.append(current_offset)
            current_offset += len(line) + 1  # +1 for newline
        full_text = "\n".join(lines)
        return full_text, line_offsets

    def _load_document(self, file_path: str) -> Tuple[str, List[int]]:
        """Load and extract text from PDF or DOCX file."""
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        suffix = file_path.suffix.lower()
        if suffix == '.pdf':
            return self._load_pdf(file_path)
        elif suffix == '.docx':
            return self._load_docx(file_path)
        else:
            raise ValueError("Unsupported file format. Only PDF and DOCX files are supported.")

    def _load_pdf(self, file_path: Path) -> Tuple[str, List[int]]:
        """Extract text from PDF preserving structure."""
        text_lines = []
        with pdfplumber.open(file_path) as pdf:
            for page_number, page in enumerate(pdf.pages, start=1):
                try:
                    # First, try extracting words
                    words = page.extract_words()
                    if words:
                        # Sort words by vertical (top) then horizontal (x0) positions
                        words.sort(key=lambda w: (w['top'], w['x0']))
                        current_line = []
                        current_top = words[0]['top']
                        for word in words:
                            # If vertical gap indicates a new line, save current_line and start a new one
                            if abs(word['top'] - current_top) > 3:
                                line_text = ' '.join(current_line)
                                if not re.search(r'(?:Source|Page|For Official Use Only|^\d+$)', line_text):
                                    text_lines.append(line_text)
                                current_line = []
                                current_top = word['top']
                            current_line.append(word['text'])
                        # Append any remaining words as the last line for this page
                        if current_line:
                            line_text = ' '.join(current_line)
                            if not re.search(r'(?:Source|Page|For Official Use Only|^\d+$)', line_text):
                                text_lines.append(line_text)
                    else:
                        # If no words were extracted, fall back to raw text extraction.
                        logger.warning(f"No words extracted on page {page_number}. Falling back to extract_text().")
                        page_text = page.extract_text()
                        if page_text:
                            for line in page_text.splitlines():
                                stripped_line = line.strip()
                                if stripped_line and not re.search(r'(?:Source|Page|For Official Use Only|^\d+$)', stripped_line):
                                    text_lines.append(stripped_line)
                        else:
                            logger.warning(f"No text found on page {page_number} with extract_text() either.")
                except Exception as e:
                    logger.error(f"Error processing PDF page {page_number}: {e}")
                    continue
        return self._join_lines(text_lines)

    def _load_docx(self, file_path: Path) -> Tuple[str, List[int]]:
        """Extract text from DOCX preserving structure."""
        doc = Document(file_path)
        text_lines = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_lines.append(para.text.strip())
                
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_lines.append(row_text)
                    
        return self._join_lines(text_lines)

    def _analyze_requirement(self, sentence: str) -> Dict:
        """Analyze a sentence to determine if it's a requirement."""
        is_mandatory = any(pattern.search(sentence) for pattern in self.mandatory_patterns)
        is_informative = any(pattern.search(sentence) for pattern in self.informative_patterns)
        has_action = any(verb in sentence.lower() for verb in self.action_verbs)
        
        # Calculate confidence score
        confidence = 0.0
        if is_mandatory:
            confidence = 0.8
        elif is_informative:
            confidence = 0.6
            
        if has_action:
            confidence = min(1.0, confidence + 0.1)
            
        return {
            'is_requirement': is_mandatory or is_informative,
            'type': 'Mandatory' if is_mandatory else ('Informative' if is_informative else None),
            'confidence': confidence
        }

    def _parse_sections(self, text: str) -> List[Dict]:
        """Parse section headers and their content."""
        lines = text.split('\n')
        sections = []
        current_section = None
        current_content = []
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
                
            # Check for section headers
            for pattern in self.header_patterns:
                match = pattern.match(line)
                if match:
                    # Save previous section
                    if current_section:
                        sections.append({
                            'id': current_section,
                            'content': ' '.join(current_content),
                            'start_line': len(sections),
                            'end_line': i
                        })
                    # Start new section
                    current_section = match.group(1)
                    current_content = [match.group(2)]
                    break
            else:
                if current_section:
                    current_content.append(line)
        
        # Add final section
        if current_section:
            sections.append({
                'id': current_section,
                'content': ' '.join(current_content),
                'start_line': len(sections),
                'end_line': len(lines)
            })
            
        return sections

    def _get_section_for_offset(self, char_offset: int, sections: List[Dict], text: str) -> Optional[str]:
        """Find the section ID for a given character offset."""
        text_before = text[:char_offset]
        line_count = text_before.count('\n')
        
        for section in sections:
            if section['start_line'] <= line_count <= section['end_line']:
                return section['id']
        return None

    def process_document(self, file_path: str) -> List[Dict]:
        """Process document and extract requirements."""
        try:
            # Load document
            logger.info(f"Processing document: {file_path}")
            text, line_offsets = self._load_document(file_path)
            
            # Parse sections
            sections = self._parse_sections(text)
            logger.info(f"Found {len(sections)} sections")
            
            # Process text with spaCy
            doc = nlp(text)
            requirements = []
            
            # Extract requirements from sentences
            for sent in doc.sents:
                # Skip short sentences
                if len(sent.text.split()) < 5:
                    continue
                    
                # Analyze the sentence
                analysis = self._analyze_requirement(sent.text)
                if not analysis['is_requirement']:
                    continue
                    
                # Find section for this sentence
                sent_section = self._get_section_for_offset(sent.start_char, sections, text)
                
                # Create requirement
                requirement = {
                    'text': sent.text.strip(),
                    'type': analysis['type'],
                    'confidence': analysis['confidence'],
                    'section_id': sent_section or 'General'
                }
                
                requirements.append(requirement)
            
            # Sort requirements
            requirements.sort(key=lambda x: (x['section_id'], -x['confidence']))
            
            logger.info(f"Extracted {len(requirements)} requirements")
            return requirements
            
        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            import traceback
            logger.error(traceback.format_exc())
            raise
