"""
SOW Analyzer Web Application
"""

import streamlit as st

# Must be the first Streamlit command
st.set_page_config(
    page_title="SOW Analyzer",
    page_icon="📄",
    layout="wide"
)

from pathlib import Path
import pandas as pd
import os
import json
import io
from dotenv import load_dotenv
import anthropic

# Load environment variables
load_dotenv()
import pdfplumber
from docx import Document
from sow_processor import SOWProcessor
from section_parser import SectionParser
from nlp_extractor import NLPExtractor

# Initialize PyTorch in a way that avoids file watcher issues
@st.cache_resource(show_spinner=False)
def get_torch():
    """Get PyTorch instance with caching to prevent reinitialization"""
    import importlib
    torch_spec = importlib.util.find_spec("torch")
    if torch_spec is None:
        st.warning("PyTorch not found")
        return None
    
    try:
        torch = importlib.import_module("torch")
        torch.set_grad_enabled(False)  # Disable gradients since we're only doing inference
        if hasattr(torch, 'jit'):
            torch.jit.disable()
        return torch
    except Exception as e:
        st.warning(f"PyTorch initialization warning: {str(e)}")
        return None

# Initialize PyTorch at startup
torch = get_torch()

def check_api_key():
    """Check if ANTHROPIC_API_KEY is set and valid"""
    try:
        api_key = os.getenv("ANTHROPIC_API_KEY")
        if not api_key:
            api_key = st.secrets["general"]["ANTHROPIC_API_KEY"]
        if not api_key or api_key == 'your-api-key-here':
            st.error('Please set your ANTHROPIC_API_KEY in environment variables or Streamlit secrets')
            return False, None
        return True, api_key
    except Exception:
        st.error('ANTHROPIC_API_KEY not found in environment variables or Streamlit secrets')
        return False, None

def init_temp_dir() -> Path:
    """Initialize temp directory."""
    if 'temp_dir' not in st.session_state:
        # Create temp directory in current working directory
        temp_dir = Path(os.getcwd()) / 'temp'
        temp_dir.mkdir(exist_ok=True, parents=True)
        st.session_state.temp_dir = temp_dir
    return st.session_state.temp_dir

def cleanup_temp_files():
    """Clean up any temporary files."""
    if 'temp_dir' in st.session_state:
        temp_dir = st.session_state.temp_dir
        if temp_dir.exists():
            for temp_file in temp_dir.glob('*'):
                try:
                    temp_file.unlink()
                except Exception as e:
                    st.error(f"Error deleting temp file {temp_file}: {str(e)}")

def main():
    st.title("SOW Analyzer")
    st.write("Extract and analyze requirements from Statement of Work (SOW) documents")
    
    # Check API key but don't stop execution
    has_api, api_key = check_api_key()
    
    # Initialize session state
    if 'initialized' not in st.session_state:
        st.session_state.processor = SOWProcessor()
        st.session_state.section_parser = SectionParser()
        st.session_state.nlp_extractor = NLPExtractor()
        if has_api:
            st.session_state.client = anthropic.Anthropic(api_key=api_key)
        st.session_state.requirements = None
        st.session_state.proposal_text = None
        st.session_state.analysis_results = None
        st.session_state.progress = 0
        st.session_state.initialized = True
    
    # Create columns for document upload
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("SOW Document")
        sow_file = st.file_uploader(
            "Upload SOW Document (PDF/DOCX)",
            type=['pdf', 'docx'],
            key='sow_uploader'
        )
    
    with col2:
        st.subheader("Proposal Document (Optional)")
        proposal_file = st.file_uploader(
            "Upload Proposal Document (Optional) (PDF/DOCX)",
            type=['pdf', 'docx'],
            key='proposal_uploader'
        )
        
        # Add checkbox for proposal matching
        perform_matching = st.checkbox(
            "Perform proposal matching analysis",
            value=False,
            help="Enable to analyze requirements against the proposal document",
            key="perform_matching"
        )
        
        # Clear analysis results when checkbox state changes
        if "last_matching_state" not in st.session_state:
            st.session_state.last_matching_state = perform_matching
        if st.session_state.last_matching_state != perform_matching:
            st.session_state.analysis_results = None
            st.session_state.last_matching_state = perform_matching
        
        if proposal_file:
            try:
                # Save uploaded file temporarily
                temp_dir = init_temp_dir()
                file_ext = Path(proposal_file.name).suffix.lower()
                temp_path = temp_dir / f"temp_proposal{file_ext}"
                if temp_path.exists():
                    temp_path.unlink()
                temp_path.write_bytes(proposal_file.getvalue())
                temp_path.chmod(0o644)  # Set read permissions
                
                try:
                    # Extract text from proposal
                    with st.spinner("Processing proposal document..."):
                        progress_bar = st.progress(0)
                        st.info("Starting document processing...")
                        proposal_text = ""
                        if temp_path.suffix.lower() == '.pdf':
                            with pdfplumber.open(temp_path) as pdf:
                                total_pages = len(pdf.pages)
                                st.info(f"Processing {total_pages} pages...")
                                for page_num, page in enumerate(pdf.pages, 1):
                                    progress_bar.progress(page_num / total_pages)
                                    st.info(f"Processing page {page_num}/{total_pages}")
                                    try:
                                        # Try layout-aware extraction first
                                        text = page.extract_text(
                                            layout=True,
                                            x_tolerance=3,
                                            y_tolerance=3,
                                            keep_blank_chars=False
                                        )
                                        if not text:  # If layout extraction fails, try basic
                                            text = page.extract_text(layout=False)
                                        if text:
                                            proposal_text += text + "\n"
                                    except Exception as e:
                                        # Log warning but continue processing
                                        st.info(f"Note: Using basic extraction for page {page_num}")
                                        text = page.extract_text(layout=False)
                                        if text:
                                            proposal_text += text + "\n"
                                        continue
                        else:  # .docx
                            try:
                                doc = Document(temp_path)
                                total_elements = len(doc.paragraphs) + len(doc.tables)
                                current_element = 0
                                st.info(f"Processing document with {total_elements} elements...")
                                
                                # Extract text from paragraphs and tables
                                for para in doc.paragraphs:
                                    current_element += 1
                                    progress_bar.progress(current_element / total_elements)
                                    if para.text.strip():
                                        proposal_text += para.text + "\n"
                                for table in doc.tables:
                                    current_element += 1
                                    progress_bar.progress(current_element / total_elements)
                                    for row in table.rows:
                                        row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                                        if row_text:
                                            proposal_text += row_text + "\n"
                            except Exception as e:
                                st.error(f"Error processing DOCX: {str(e)}")
                                raise
                        
                        if not proposal_text.strip():
                            raise ValueError("No text could be extracted from the document")
                        
                        st.session_state.proposal_text = proposal_text
                        progress_bar.progress(1.0)
                        st.success(f"Proposal document processed successfully! Extracted {len(proposal_text.split())} words.")
                
                except Exception as e:
                    st.error(f"Error processing proposal: {str(e)}")
                    st.session_state.proposal_text = None
                
                finally:
                    # Cleanup temp file
                    if temp_path.exists():
                        temp_path.unlink()
                        
            except Exception as e:
                st.error(f"Error handling proposal file: {str(e)}")
    
    if sow_file:
        try:
            # Save uploaded file temporarily
            temp_dir = init_temp_dir()
            file_ext = Path(sow_file.name).suffix.lower()
            temp_path = temp_dir / f"temp_sow{file_ext}"
            if temp_path.exists():
                temp_path.unlink()
            temp_path.write_bytes(sow_file.getvalue())
            temp_path.chmod(0o644)  # Set read permissions
            
            try:
                # Process SOW document
                with st.spinner("Processing SOW document..."):
                    # Extract text
                    text = st.session_state.processor._load_document(temp_path)
                    
                    # Parse sections
                    sections = st.session_state.section_parser.parse_sections(text)
                    
                    # Extract requirements from each section
                    requirements = []
                    for section in sections:
                        section_reqs = st.session_state.nlp_extractor.extract_requirements(
                            '\n'.join(section.content),
                            section.id,
                            section.title
                        )
                        requirements.extend([{
                            'section_id': req.section_id,
                            'section_title': req.section_title,
                            'text': req.text,
                            'type': req.type,
                            'confidence': req.confidence,
                            'entities': req.entities
                        } for req in section_reqs])
                    
                    st.session_state.requirements = requirements
                    st.session_state.progress = 1.0
                
                # Display initial results
                st.success(f"Found {len(requirements)} requirements!")
                
                # Display requirements by type
                req_types = {}
                for req in requirements:
                    req_type = req['type']
                    req_types[req_type] = req_types.get(req_type, 0) + 1
                
                # Create columns for metrics
                cols = st.columns(len(req_types))
                for i, (req_type, count) in enumerate(req_types.items()):
                    with cols[i]:
                        st.metric(f"{req_type} Requirements", count)
                
                # Display requirements table
                st.subheader("Requirements Analysis")
                
                # Create DataFrame
                df = pd.DataFrame(requirements)
                
                # Add entity columns
                all_entity_types = set()
                for req in requirements:
                    all_entity_types.update(req['entities'].keys())
                
                for ent_type in all_entity_types:
                    df[f'entities_{ent_type}'] = df['entities'].apply(
                        lambda x: ', '.join(x.get(ent_type, [])) if ent_type in x else ''
                    )
                
                # Drop the original entities column
                df = df.drop('entities', axis=1)
                
                # Configure column display
                column_config = {
                    "section_id": "Section",
                    "section_title": "Section Title",
                    "text": "Requirement",
                    "type": "Type",
                    "confidence": st.column_config.NumberColumn(
                        "Confidence",
                        help="Confidence score (0-1)",
                        format="%.2f"
                    )
                }
                
                # Add entity column configs
                for ent_type in all_entity_types:
                    column_config[f'entities_{ent_type}'] = st.column_config.TextColumn(
                        f"{ent_type} Entities",
                        help=f"Extracted {ent_type} entities"
                    )
                
                # Only perform proposal matching if enabled and proposal is available
                if not proposal_file:
                    if perform_matching:
                        st.info("Upload a proposal document to analyze requirements compliance")
                elif st.session_state.proposal_text and perform_matching and not st.session_state.analysis_results:
                    # Start analysis
                    with st.spinner("Analyzing requirements against proposal..."):
                            results = []
                            total_reqs = len(requirements)
                            progress_bar = st.progress(0)
                            analysis_status = st.empty()
                            
                            # Create a placeholder for the live results table
                            results_table = st.empty()
                            
                            for i, req in enumerate(requirements, 1):
                                progress = i / total_reqs
                                progress_bar.progress(progress)
                                analysis_status.write(f"Analyzing requirement {i} of {total_reqs} ({int(progress * 100)}%)")
                                
                                prompt = f"""You are a requirements analysis assistant. Find the section in the proposal that best matches this SOW requirement. Look for section numbers in the format X.X.X or X.X at the start of paragraphs. Return ONLY a JSON object with no additional text.
SOW Requirement (Section {req['section_id']}):
{req['text']}
Proposal Text:
{st.session_state.proposal_text}
Return a JSON object with these exact fields:
{{
    "matched_section": "The exact section number found in the proposal (e.g., 1.1.1, 2.3, etc.). If no numbered section matches, return 'N/A'",
    "matched_text": "The exact text from that section that addresses this requirement",
    "compliance": "Fully Compliant" or "Partially Compliant" or "Not Addressed",
    "confidence": A number between 0 and 1 indicating match confidence,
    "suggestions": ["List", "of", "improvement", "suggestions"]
}}"""
                                
                                if has_api:
                                    response = st.session_state.client.messages.create(
                                        model="claude-3-opus-20240229",
                                        max_tokens=1000,
                                        messages=[{
                                            "role": "user",
                                            "content": prompt
                                        }]
                                    )
                                else:
                                    st.error("Cannot analyze proposal without ANTHROPIC_API_KEY")
                                    break
                                
                                try:
                                    # Extract JSON from response
                                    response_text = response.content[0].text.strip()
                                    
                                    # If no proposal text, return a default analysis
                                    if not st.session_state.proposal_text or not st.session_state.proposal_text.strip():
                                        analysis = {
                                            'matched_section': 'N/A',
                                            'matched_text': '',
                                            'compliance': 'Not Addressed',
                                            'confidence': 0.0,
                                            'suggestions': ['Upload a proposal document to analyze']
                                        }
                                    else:
                                        # Remove any non-JSON text before or after
                                        json_start = response_text.find('{')
                                        json_end = response_text.rfind('}') + 1
                                        if json_start >= 0 and json_end > json_start:
                                            json_str = response_text[json_start:json_end]
                                            analysis = json.loads(json_str)
                                        else:
                                            # If no JSON found, create a default analysis
                                            analysis = {
                                                'matched_section': 'Error',
                                                'matched_text': '',
                                                'compliance': 'Not Analyzed',
                                                'confidence': 0.0,
                                                'suggestions': ['Try uploading the proposal document again']
                                            }
                                    
                                    result = {
                                        **req,
                                        'matched_section': analysis.get('matched_section'),
                                        'matched_text': analysis.get('matched_text', ''),
                                        'compliance': analysis.get('compliance', 'Not Analyzed'),
                                        'match_confidence': analysis.get('confidence', 0.0),
                                        'suggestions': analysis.get('suggestions', [])
                                    }
                                    results.append(result)
                                    
                                    # Update the live results table
                                    results_df = pd.DataFrame(results)
                                    results_table.dataframe(
                                        results_df,
                                        column_config={
                                            "section_id": "SOW Section",
                                            "text": "Requirement",
                                            "type": "Type",
                                            "confidence": st.column_config.NumberColumn(
                                                "Extraction Confidence",
                                                help="Confidence in requirement extraction (0-1)",
                                                format="%.2f"
                                            ),
                                            "matched_section": "Matched Proposal Section",
                                            "matched_text": "Matched Proposal Text",
                                            "compliance": "Compliance Status",
                                            "match_confidence": st.column_config.NumberColumn(
                                                "Match Confidence",
                                                help="Confidence in proposal match (0-1)",
                                                format="%.2f"
                                            ),
                                            "suggestions": "Improvement Suggestions"
                                        },
                                        hide_index=True
                                    )
                                except Exception as e:
                                    st.error(f"Error parsing analysis for requirement {req['section_id']}: {str(e)}")
                                    st.error(f"Raw response: {response_text}")
                                    results.append({
                                        **req,
                                        'matched_text': '',
                                        'compliance': 'Error',
                                        'match_confidence': 0.0,
                                        'suggestions': []
                                    })
                            
                            progress_bar.progress(1.0)
                            analysis_status.write("Analysis complete!")
                            st.session_state.analysis_results = results
                    
                    # Display analysis results
                    df = pd.DataFrame(st.session_state.analysis_results)
                    st.dataframe(
                        df,
                        column_config={
                            "section_id": "SOW Section",
                            "text": "Requirement",
                            "type": "Type",
                            "confidence": st.column_config.NumberColumn(
                                "Extraction Confidence",
                                help="Confidence in requirement extraction (0-1)",
                                format="%.2f"
                            ),
                            "matched_section": "Matched Proposal Section",
                            "matched_text": "Matched Proposal Text",
                            "compliance": "Compliance Status",
                            "match_confidence": st.column_config.NumberColumn(
                                "Match Confidence",
                                help="Confidence in proposal match (0-1)",
                                format="%.2f"
                            ),
                            "suggestions": "Improvement Suggestions"
                        },
                        hide_index=True
                    )
                else:
                    # Display basic requirements table
                    st.dataframe(
                        df,
                        column_config=column_config,
                        hide_index=True
                    )
                
                # Export options
                col1, col2, col3 = st.columns(3)
                with col1:
                    # Convert requirements to CSV
                    csv_data = df.to_csv(index=False)
                    st.download_button(
                        "Download Requirements (CSV)",
                        csv_data,
                        "sow_requirements.csv",
                        "text/csv",
                        key='download-csv'
                    )
                
                with col2:
                    # Convert requirements to Excel
                    buffer = io.BytesIO()
                    with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                        df.to_excel(writer, index=False)
                    st.download_button(
                        "Download Requirements (Excel)",
                        buffer.getvalue(),
                        "sow_requirements.xlsx",
                        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        key='download-excel'
                    )
                
                with col3:
                    if st.session_state.analysis_results:
                        # Convert analysis results to Excel
                        analysis_buffer = io.BytesIO()
                        with pd.ExcelWriter(analysis_buffer, engine='openpyxl') as writer:
                            pd.DataFrame(st.session_state.analysis_results).to_excel(writer, index=False)
                        st.download_button(
                            "Download Analysis (Excel)",
                            analysis_buffer.getvalue(),
                            "sow_analysis_results.xlsx",
                            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            key='download-analysis'
                        )
                
            except Exception as e:
                st.error(f"Error processing document: {str(e)}")
                st.session_state.requirements = None
                st.session_state.progress = 0
            
            finally:
                # Cleanup temp file
                if temp_path.exists():
                    temp_path.unlink()
                
        except Exception as e:
            st.error(f"Error handling file: {str(e)}")

if __name__ == "__main__":
    main()
